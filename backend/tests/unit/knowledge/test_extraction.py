"""Tests for text extraction from PDF, DOCX, and TXT files."""

from __future__ import annotations

import pytest

from yourai.knowledge.extraction import ExtractionResult, Section, extract_text


class TestExtractTxt:
    """Tests for plain text extraction."""

    def test_utf8_text(self):
        data = b"Hello, world! This is a test."
        result = extract_text(data, "text/plain", "test.txt")
        assert "Hello, world!" in result.text
        assert result.strategy == "txt_utf-8"
        assert len(result.sections) == 1

    def test_latin1_text(self):
        data = "Héllo, wörld!".encode("latin-1")
        result = extract_text(data, "text/plain", "test.txt")
        assert "Héllo" in result.text

    def test_empty_text(self):
        data = b""
        result = extract_text(data, "text/plain", "test.txt")
        assert result.text == ""


class TestExtractPdf:
    """Tests for PDF extraction using PyMuPDF."""

    def test_simple_pdf(self):
        """Test extraction from a minimal valid PDF."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test PDF Content", fontsize=12)
        pdf_bytes = doc.write()
        doc.close()

        result = extract_text(pdf_bytes, "application/pdf", "test.pdf")
        assert "Test PDF Content" in result.text
        assert result.strategy == "pdf_pymupdf"

    def test_pdf_with_headings(self):
        """Test that larger font text is detected as headings."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        # Large text = heading
        page.insert_text((72, 72), "Main Title", fontsize=24)
        # Normal text = body
        page.insert_text((72, 120), "Body text goes here.", fontsize=11)
        page.insert_text((72, 140), "More body text.", fontsize=11)
        pdf_bytes = doc.write()
        doc.close()

        result = extract_text(pdf_bytes, "application/pdf", "test.pdf")
        assert "Main Title" in result.text
        assert "Body text" in result.text
        # Should have at least one section with a heading
        headings = [s for s in result.sections if s.heading is not None]
        assert len(headings) >= 1

    def test_empty_pdf(self):
        """Test extraction from an empty PDF."""
        import fitz

        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.write()
        doc.close()

        result = extract_text(pdf_bytes, "application/pdf", "test.pdf")
        assert result.text == ""


class TestExtractDocx:
    """Tests for DOCX extraction using python-docx."""

    def test_simple_docx(self):
        """Test extraction from a simple DOCX."""
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Test paragraph content.")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        result = extract_text(
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "test.docx",
        )
        assert "Test paragraph content" in result.text
        assert result.strategy == "docx_python_docx"

    def test_docx_with_headings(self):
        """Test that heading styles are detected."""
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("Chapter content here.")
        doc.add_heading("Chapter 2", level=1)
        doc.add_paragraph("More content here.")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        result = extract_text(
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "test.docx",
        )
        assert len(result.sections) >= 2
        headings = [s for s in result.sections if s.heading is not None]
        assert len(headings) >= 2
        assert "Chapter 1" in headings[0].heading


class TestExtractUnsupported:
    """Test unsupported MIME types."""

    def test_unsupported_raises(self):
        with pytest.raises(ValueError, match="Unsupported MIME type"):
            extract_text(b"data", "image/png", "test.png")


class TestExtractionResult:
    """Tests for the ExtractionResult dataclass."""

    def test_creation(self):
        result = ExtractionResult(text="hello", strategy="test")
        assert result.text == "hello"
        assert result.sections == []

    def test_with_sections(self):
        sections = [Section(heading="H1", content="Content", level=1)]
        result = ExtractionResult(text="hello", strategy="test", sections=sections)
        assert len(result.sections) == 1
